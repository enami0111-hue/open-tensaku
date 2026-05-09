"""テキスト本文を新規.docxに変換し、AI添削コメントを add_comment() で注入する。

MVPの単純化方針：
- 入力テキストの各行 → 1段落 → 1run
- 添削コメントは `target_text` を含む段落全体にアンカーする（精細な範囲指定はしない）
- これによりrun分割の複雑性を回避し、Word/Pages/Google Docsでの互換性も最大化
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

AUTHOR = "オープン添削"
INITIALS = "OT"


@dataclass
class BuildResult:
    docx_bytes: bytes
    matched: int
    unmatched: list[dict[str, Any]]


GRADE_COLORS = {
    "A": RGBColor(0x0F, 0x8B, 0x5C),
    "B": RGBColor(0x0F, 0x6F, 0xCF),
    "C": RGBColor(0xFF, 0xA7, 0x03),
    "D": RGBColor(0xC8, 0x10, 0x2E),
}


def build_corrected_docx(
    *,
    university_name: str,
    faculty_name: str,
    course_name: str,
    question_title: str,
    question_body: str,
    question_year: int,
    question_status: str,
    student_body: str,
    correction: dict[str, Any],
) -> BuildResult:
    doc = Document()
    _set_default_font(doc)

    summary = correction.get("summary", {})
    comments = correction.get("comments", [])
    good_count = summary.get("good_count", sum(1 for c in comments if c.get("type") == "good"))
    more_count = summary.get("more_count", sum(1 for c in comments if c.get("type") == "more"))
    diagnosis = correction.get("diagnosis", {}) or {}
    grade = diagnosis.get("grade") or summary.get("grade")
    headline = diagnosis.get("headline", "")
    # opening は新スキーマで diagnosis.opening、後方互換で top-level の opening も参照
    opening = (diagnosis.get("opening") or correction.get("opening") or "").strip()
    section_summary = correction.get("section_summary") or []
    next_actions = correction.get("next_actions") or []

    title = doc.add_heading("【オープン添削】添削結果", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)

    meta = doc.add_paragraph()
    meta.add_run(f"志望先：{university_name} {faculty_name}（{course_name}）\n").bold = True
    year_label = "当年度" if question_status == "official" else "前年度（参考）"
    meta.add_run(f"設問：{question_title} ／ {year_label}：{question_year}年度\n")
    meta.add_run(f"添削サマリー：Good {good_count}件・More {more_count}件")

    # 全体評価（grade + headline）
    if grade or headline:
        doc.add_heading("全体評価", level=1)
        diag_para = doc.add_paragraph()
        if grade:
            r = diag_para.add_run(f"【グレード {grade}】 ")
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = GRADE_COLORS.get(grade, RGBColor(0x07, 0x37, 0x63))
        if headline:
            diag_para.add_run(headline).bold = True

    # 開幕の労い
    if opening:
        op_para = doc.add_paragraph(opening)
        op_para.style = doc.styles["Intense Quote"]

    doc.add_heading("設問", level=1)
    q_para = doc.add_paragraph(question_body.strip())
    q_para.style = doc.styles["Intense Quote"]

    doc.add_heading("あなたの志望理由書（添削コメント付き）", level=1)

    body_paragraphs: list[Any] = []
    for line in student_body.splitlines() or [""]:
        if not line.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line)
        body_paragraphs.append(p)

    matched = 0
    unmatched: list[dict[str, Any]] = []
    for c in comments:
        target = (c.get("target_text") or "").strip()
        comment_body = (c.get("comment") or "").strip()
        if not target or not comment_body:
            unmatched.append({**c, "_reason": "empty_target_or_comment"})
            continue

        prefix = "【Good】" if c.get("type") == "good" else "【More】"
        cat = c.get("category", "")
        section = c.get("section", "")
        meta_bits = []
        if section:
            meta_bits.append(f"#{section}")
        if cat:
            meta_bits.append(f"観点：{cat}")
        meta_suffix = f"（{' / '.join(meta_bits)}）" if meta_bits else ""
        full_comment = f"{prefix}{meta_suffix}\n{comment_body}"

        para = _find_paragraph_containing(body_paragraphs, target)
        if para is None:
            unmatched.append({**c, "_reason": "target_text not found in body"})
            continue

        runs = list(para.runs)
        if not runs:
            unmatched.append({**c, "_reason": "paragraph has no runs"})
            continue

        try:
            doc.add_comment(
                runs=runs,
                text=full_comment,
                author=AUTHOR,
                initials=INITIALS,
            )
            matched += 1
        except Exception as e:
            unmatched.append({**c, "_reason": f"add_comment_error: {e}"})

    # セクション別講評
    if section_summary:
        doc.add_heading("セクション別講評", level=1)
        for s in section_summary:
            sec = (s.get("section") or "全体").strip()
            txt = (s.get("comment") or "").strip()
            if not txt:
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"【{sec}】 ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x07, 0x37, 0x63)
            p.add_run(txt)

    # 次の一手（新スキーマは [{priority, action}]、旧スキーマは [str] を許容）
    if next_actions:
        doc.add_heading("次の一手（改稿の優先順位）", level=1)
        # priority昇順にソート（指定なしは最後）
        def _prio(a):
            return a.get("priority", 999) if isinstance(a, dict) else 999
        for action in sorted(next_actions, key=_prio):
            if isinstance(action, dict):
                text = (action.get("action") or "").strip()
            else:
                text = str(action).strip()
            if text:
                doc.add_paragraph(text, style="List Number")

    # 補足コメント
    if unmatched:
        doc.add_heading("補足コメント（本文への直接アンカー失敗分）", level=1)
        for c in unmatched:
            prefix = "【Good】" if c.get("type") == "good" else "【More】"
            doc.add_paragraph(
                f'{prefix} 「{c.get("target_text", "")}」について：{c.get("comment", "")}',
                style="List Bullet",
            )

    # 励まし（最後に）
    encouragement = (correction.get("encouragement") or "").strip()
    if encouragement:
        doc.add_heading("オープン添削からひとこと", level=1)
        p = doc.add_paragraph()
        run = p.add_run(encouragement)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)

    buf = io.BytesIO()
    doc.save(buf)
    return BuildResult(docx_bytes=buf.getvalue(), matched=matched, unmatched=unmatched)


def _set_default_font(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "ヒラギノ角ゴ Pro"
    style.font.size = Pt(11)


def _normalize(s: str) -> str:
    return re.sub(r"\s+", "", s)


def _find_paragraph_containing(paragraphs: list[Any], target: str):
    target_norm = _normalize(target)
    if not target_norm:
        return None
    for para in paragraphs:
        text = "".join(r.text for r in para.runs)
        if target in text:
            return para
    for para in paragraphs:
        text = "".join(r.text for r in para.runs)
        if target_norm in _normalize(text):
            return para
    return None
